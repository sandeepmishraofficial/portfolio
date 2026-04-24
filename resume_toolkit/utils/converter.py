"""
converter.py — PDF ↔ DOCX conversion utilities.
PDF → DOCX: PyMuPDF (fitz) + python-docx
DOCX → PDF: python-docx + reportlab
"""
import io


def pdf_to_docx(file_content: bytes) -> bytes:
    """Convert PDF bytes to DOCX bytes."""
    try:
        import fitz
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError as e:
        raise ImportError(f"Missing library: {e}. Run: pip install PyMuPDF python-docx")

    try:
        pdf_doc = fitz.open(stream=file_content, filetype="pdf")
        word_doc = Document()

        # Set document margins
        for section in word_doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        # Set default font
        style = word_doc.styles['Normal']
        font  = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        for page_num, page in enumerate(pdf_doc):
            if page_num > 0:
                word_doc.add_paragraph()  # Page separator

            # Extract text blocks sorted top-to-bottom, left-to-right
            blocks = page.get_text("blocks")
            blocks_sorted = sorted(blocks, key=lambda b: (round(b[1] / 20), b[0]))

            for block in blocks_sorted:
                block_text = block[4].strip()
                if not block_text:
                    continue
                for line in block_text.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    para = word_doc.add_paragraph(line)
                    run  = para.runs[0] if para.runs else para.add_run(line)
                    run.font.size = Pt(11)

        pdf_doc.close()

        output = io.BytesIO()
        word_doc.save(output)
        output.seek(0)
        return output.getvalue()

    except Exception as e:
        raise Exception(f"PDF to DOCX conversion failed: {e}")


def docx_to_pdf(file_content: bytes) -> bytes:
    """Convert DOCX bytes to PDF bytes using reportlab."""
    try:
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
    except ImportError as e:
        raise ImportError(f"Missing library: {e}. Run: pip install python-docx reportlab")

    try:
        docx_doc = Document(io.BytesIO(file_content))

        output = io.BytesIO()
        pdf = SimpleDocTemplate(
            output,
            pagesize=A4,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )

        styles = getSampleStyleSheet()

        # Custom styles
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=colors.HexColor('#1e3a5f'),
            spaceBefore=12,
            spaceAfter=4,
            fontName='Helvetica-Bold',
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10.5,
            leading=16,
            spaceAfter=4,
            fontName='Helvetica',
        )
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontSize=10.5,
            leading=16,
            leftIndent=20,
            spaceAfter=2,
            fontName='Helvetica',
        )

        story = []

        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 0.08 * inch))
                continue

            # Escape XML special chars for reportlab
            safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            # Detect heading (short, bold-ish, or ALL CAPS)
            is_heading = (
                (len(text) < 80 and text.isupper()) or
                (para.style and 'Heading' in para.style.name)
            )

            # Detect bullet
            is_bullet = text.startswith(('-', '•', '*', '·', '◦', '▪', '▸', '►'))

            if is_heading:
                story.append(Paragraph(safe_text, heading_style))
            elif is_bullet:
                story.append(Paragraph(f'• {safe_text.lstrip("-•*·◦▪▸► ").strip()}', bullet_style))
            else:
                story.append(Paragraph(safe_text, body_style))

        if not story:
            story.append(Paragraph("(Empty document)", body_style))

        pdf.build(story)
        output.seek(0)
        return output.getvalue()

    except Exception as e:
        raise Exception(f"DOCX to PDF conversion failed: {e}")
