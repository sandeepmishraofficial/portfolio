"""
extractor.py — Extract text from PDF and DOCX files.
Uses PyMuPDF for PDFs and python-docx for DOCX files.
"""
import io


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract plain text from PDF bytes using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is not installed. Run: pip install PyMuPDF")

    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())
        doc.close()
        return "\n".join(pages_text).strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract plain text from DOCX bytes using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    try:
        doc = Document(io.BytesIO(file_content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return "\n".join(paragraphs)
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {e}")


def extract_text(file_content: bytes, filename: str) -> str:
    """Auto-detect file type from filename and extract text."""
    name_lower = filename.lower()
    if name_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_content)
    elif name_lower.endswith('.docx'):
        return extract_text_from_docx(file_content)
    elif name_lower.endswith('.doc'):
        raise ValueError("Legacy .doc format is not supported. Please save as .docx")
    else:
        raise ValueError(f"Unsupported file type '{filename}'. Please upload a PDF or DOCX.")
